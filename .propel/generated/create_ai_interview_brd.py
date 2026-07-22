from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "AI_Interview_Platform_BRD.docx"
PDF_PATH = ROOT / "AI_Interview_Platform_BRD.pdf"


def cell_fill(cell, fill):
    table_cell_props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    table_cell_props.append(shading)


def get_document_model():
    today = date.today().isoformat()
    return {
        "title": "AI Interview Platform",
        "subtitle": "Business Requirements Document",
        "tagline": "Architect perspective covering business, UI, engineering, security, and testing requirements",
        "meta": [
            ["Document Version", "1.0"],
            ["Prepared For", "AI Interview Platform stakeholders"],
            ["Prepared Date", today],
            ["Status", "Draft for business, architecture, and delivery review"],
            ["Reference Style", "Structured in the same enterprise BRD pattern as the uploaded ASG-Vasan reference"],
        ],
        "sections": [
            {
                "heading": "1. DOCUMENT CONTROL",
                "items": [
                    {"type": "subheading", "text": "1.1 Version History"},
                    {
                        "type": "table",
                        "headers": ["Version", "Date", "Author", "Change Summary", "Status"],
                        "rows": [["1.0", today, "GitHub Copilot", "Initial architect-level BRD for AI Interview Platform", "Draft"]],
                    },
                    {"type": "subheading", "text": "1.2 Review and Approval Status"},
                    {
                        "type": "table",
                        "headers": ["Role", "Owner", "Status", "Comments"],
                        "rows": [
                            ["Business Sponsor", "TBD", "Pending", "Confirm scope and ROI assumptions"],
                            ["Product Owner", "TBD", "Pending", "Confirm MVP and release sequencing"],
                            ["Solution Architect", "TBD", "Pending", "Confirm architecture and integration boundaries"],
                            ["Security and Compliance", "TBD", "Pending", "Confirm consent, retention, and audit rules"],
                        ],
                    },
                    {"type": "subheading", "text": "1.3 Document Purpose"},
                    {
                        "type": "paragraph",
                        "text": "This Business Requirements Document defines the end-to-end requirements for the AI Interview Platform. It is intended to guide discovery, architecture, UI design, implementation planning, testing, compliance validation, and rollout governance. The document is written from an architect perspective, which means it captures not only business intent but also the structural considerations needed to build the platform safely and at scale.",
                    },
                ],
            },
            {
                "heading": "2. EXECUTIVE SUMMARY",
                "items": [
                    {"type": "subheading", "text": "2.1 Organisation Background"},
                    {
                        "type": "paragraph",
                        "text": "The AI Interview Platform is a SaaS application for talent acquisition teams that need faster and more consistent candidate screening. The product supports asynchronous and live interviews, AI-generated question sets, transcription, summary generation, structured scorecards, and recruiter decision workflows. The platform is designed to augment human decision-making rather than replace it.",
                    },
                    {"type": "subheading", "text": "2.2 Business Problem"},
                    {
                        "type": "bullets",
                        "items": [
                            "Recruiters spend too much time coordinating interviews, repeating early-stage questions, and manually writing notes.",
                            "Interview quality varies by interviewer, which creates inconsistent candidate evaluation and weak comparability across applicants.",
                            "Candidates experience fragmented communication across email, calendar, and video tools, leading to drop-off and poor experience.",
                            "Leadership lacks real-time operational visibility into screening throughput, delay sources, and reviewer productivity.",
                        ],
                    },
                    {"type": "subheading", "text": "2.3 Proposed Solution"},
                    {
                        "type": "paragraph",
                        "text": "The proposed solution centralizes the screening lifecycle into one governed platform. Recruiters can configure role-based interview templates, invite candidates, collect responses, and review AI-generated summaries with evidence-linked transcripts. Hiring teams gain a standardized scorecard workflow, while administrators gain auditability, policy controls, and reporting visibility.",
                    },
                    {"type": "subheading", "text": "2.4 Strategic Value and ROI"},
                    {
                        "type": "bullets",
                        "items": [
                            "Reduce recruiter screening effort by at least 50 percent.",
                            "Improve completion rate of first-round interviews to at least 85 percent.",
                            "Shorten turnaround from candidate completion to reviewer action to under one business day.",
                            "Increase consistency of scoring and evidence capture across teams and regions.",
                            "Provide a compliance-ready evidence trail for consent, scoring overrides, and hiring decisions.",
                        ],
                    },
                ],
            },
            {
                "heading": "3. PROJECT SCOPE",
                "items": [
                    {"type": "subheading", "text": "3.1 In-Scope Modules"},
                    {
                        "type": "bullets",
                        "items": [
                            "Tenant onboarding and organization configuration",
                            "Recruiter, hiring manager, and administrator user management",
                            "Job intake and interview template configuration",
                            "Candidate invitation, authentication, and consent capture",
                            "Asynchronous and live interview session management",
                            "Audio or video recording, transcription, and AI summarization",
                            "Scorecards, reviewer workflow, and decision logging",
                            "Notifications, reminders, and escalation alerts",
                            "Reporting, analytics, and operational dashboards",
                            "ATS, calendar, identity provider, email, SMS, and AI vendor integrations",
                        ],
                    },
                    {"type": "subheading", "text": "3.2 Out-of-Scope Items"},
                    {
                        "type": "bullets",
                        "items": [
                            "Full ATS replacement",
                            "Offer management, payroll, and compensation workflows",
                            "Background verification and onboarding paperwork",
                            "Autonomous hiring decisions without explicit human review",
                            "Custom model training pipelines for v1 unless separately funded",
                        ],
                    },
                    {"type": "subheading", "text": "3.3 Assumptions"},
                    {
                        "type": "bullets",
                        "items": [
                            "Each interview flow will be tied to a role, job family, or competency blueprint.",
                            "At least one approved speech-to-text and one approved LLM provider will be available for production use.",
                            "Legal and compliance stakeholders will define candidate consent wording and jurisdictional policy settings.",
                            "The initial ATS integration will focus on candidate sync, requisition metadata, and decision status updates.",
                        ],
                    },
                ],
            },
            {
                "heading": "4. STAKEHOLDERS",
                "items": [
                    {"type": "subheading", "text": "4.1 Primary Stakeholders"},
                    {
                        "type": "table",
                        "headers": ["Stakeholder", "Responsibility", "Primary Need", "Value Expected"],
                        "rows": [
                            ["Recruiter", "Owns screening workflow", "Move candidates through early funnel faster", "Lower effort and faster decisions"],
                            ["Hiring Manager", "Reviews and approves candidates", "See structured evidence before interviews progress", "More reliable candidate comparison"],
                            ["Candidate", "Participates in the interview", "Simple, transparent, mobile-friendly experience", "Less friction and clearer expectations"],
                            ["Administrator", "Configures system settings", "Manage access, templates, and policy controls", "Stable and governed operations"],
                        ],
                    },
                    {"type": "subheading", "text": "4.2 Management Stakeholders"},
                    {
                        "type": "bullets",
                        "items": [
                            "Head of Talent Acquisition",
                            "HR Operations Lead",
                            "Regional Hiring Directors",
                            "Product Sponsor",
                            "Delivery Manager",
                        ],
                    },
                    {"type": "subheading", "text": "4.3 Secondary and External Stakeholders"},
                    {
                        "type": "bullets",
                        "items": [
                            "Security and compliance teams",
                            "Legal and privacy teams",
                            "ATS vendor or integration owner",
                            "Identity provider administrator",
                            "Email and SMS service providers",
                            "Speech-to-text and LLM providers",
                        ],
                    },
                ],
            },
            {
                "heading": "5. BUSINESS OBJECTIVES AND SUCCESS METRICS",
                "items": [
                    {"type": "subheading", "text": "5.1 Strategic Objectives"},
                    {
                        "type": "bullets",
                        "items": [
                            "Standardize interview quality across functions, interviewers, and geographies.",
                            "Improve time-to-shortlist without increasing recruiter headcount.",
                            "Reduce subjectivity by aligning interviewer behavior to structured scorecards.",
                            "Increase auditability of screening and reviewer actions.",
                        ],
                    },
                    {"type": "subheading", "text": "5.2 Success Metrics"},
                    {
                        "type": "table",
                        "headers": ["Metric", "Target", "Measurement Method"],
                        "rows": [
                            ["Recruiter screening effort", "50 percent reduction", "Baseline vs post-go-live time study"],
                            ["Candidate completion rate", "85 percent or better", "Interview completion analytics"],
                            ["Summary turnaround", "Under 5 minutes", "Workflow processing telemetry"],
                            ["Reviewer turnaround", "Under 1 business day", "Decision cycle reporting"],
                            ["Audit completeness", "100 percent of critical actions logged", "Quarterly audit review"],
                        ],
                    },
                ],
            },
            {
                "heading": "6. CURRENT STATE (AS-IS) AND PAIN POINTS",
                "items": [
                    {"type": "subheading", "text": "6.1 Process Pain Points"},
                    {
                        "type": "bullets",
                        "items": [
                            "Recruiters coordinate across too many disconnected systems.",
                            "Candidates receive inconsistent instructions and reminders.",
                            "Interview records are spread across notes, email, spreadsheets, and meeting recordings.",
                            "Managers often re-run early screening questions because prior evidence is unreliable or unavailable.",
                        ],
                    },
                    {"type": "subheading", "text": "6.2 Technology and Governance Gaps"},
                    {
                        "type": "bullets",
                        "items": [
                            "No common data model for interview sessions, scorecards, and transcripts.",
                            "No consistent access control model across candidate and reviewer data.",
                            "Limited monitoring of candidate drop-off, queue time, or vendor failures.",
                            "Inconsistent documentation of scoring rationale and decision overrides.",
                        ],
                    },
                ],
            },
            {
                "heading": "7. FUTURE STATE (TO-BE)",
                "items": [
                    {"type": "subheading", "text": "7.1 Recruiter Journey"},
                    {
                        "type": "numbered",
                        "items": [
                            "Recruiter selects a role or requisition and applies a governed interview template.",
                            "Platform generates or suggests questions aligned to required competencies.",
                            "Recruiter schedules or triggers interview invitation.",
                            "Candidate completes the interview and the platform produces transcript and summary artifacts.",
                            "Recruiter and hiring manager review the evidence and log a disposition.",
                        ],
                    },
                    {"type": "subheading", "text": "7.2 Candidate Journey"},
                    {
                        "type": "numbered",
                        "items": [
                            "Candidate receives invite with job context and expectations.",
                            "Candidate authenticates and confirms consent.",
                            "Candidate completes interview on web or mobile browser with progress visibility.",
                            "Candidate receives completion acknowledgment and next-step expectations.",
                        ],
                    },
                    {"type": "subheading", "text": "7.3 Target Outcome Improvements"},
                    {
                        "type": "bullets",
                        "items": [
                            "One system of record for screening-stage activity",
                            "Template-based governance with consistent scoring evidence",
                            "Asynchronous processing for transcription and summary generation",
                            "Operational dashboards for throughput, quality, and compliance",
                        ],
                    },
                ],
            },
            {
                "heading": "8. FUNCTIONAL REQUIREMENTS",
                "items": [
                    {"type": "subheading", "text": "8.1 Candidate and Access Module (FR-CAN)"},
                    {
                        "type": "table",
                        "headers": ["ID", "Requirement", "Priority", "Architect Note"],
                        "rows": [
                            ["FR-CAN-01", "System shall support candidate login by secure magic link or OTP.", "High", "Avoid password management for short-lived candidate journeys"],
                            ["FR-CAN-02", "System shall capture explicit consent before recording or AI processing begins.", "High", "Consent must be stored as a separate auditable record"],
                            ["FR-CAN-03", "System shall support resume or profile upload and validation.", "Medium", "Store upload references outside transactional row payloads"],
                        ],
                    },
                    {"type": "subheading", "text": "8.2 Job Intake and Template Management (FR-TPL)"},
                    {
                        "type": "table",
                        "headers": ["ID", "Requirement", "Priority", "Architect Note"],
                        "rows": [
                            ["FR-TPL-01", "System shall create reusable templates by role, level, and competency area.", "High", "Templates should be versioned and immutable after publication"],
                            ["FR-TPL-02", "System shall allow question pools, time limits, and scoring rubric configuration.", "High", "Keep rubric as structured data, not free text only"],
                            ["FR-TPL-03", "System shall allow AI-assisted question generation from job description input.", "High", "Generated content requires review workflow before production use"],
                        ],
                    },
                    {"type": "subheading", "text": "8.3 Scheduling and Invitations (FR-SCH)"},
                    {
                        "type": "table",
                        "headers": ["ID", "Requirement", "Priority", "Architect Note"],
                        "rows": [
                            ["FR-SCH-01", "System shall send invitations by email and optional SMS.", "High", "Track delivery and bounce status"],
                            ["FR-SCH-02", "System shall support time-window validity and reminder schedules.", "High", "Use event-driven reminders rather than synchronous cron-heavy logic"],
                            ["FR-SCH-03", "System shall integrate with calendar providers for live interviews.", "Medium", "Live sessions may be phased after asynchronous MVP"],
                        ],
                    },
                    {"type": "subheading", "text": "8.4 Interview Delivery (FR-INT)"},
                    {
                        "type": "table",
                        "headers": ["ID", "Requirement", "Priority", "Architect Note"],
                        "rows": [
                            ["FR-INT-01", "System shall support asynchronous recorded interviews.", "High", "MVP anchor capability"],
                            ["FR-INT-02", "System shall support live interviews with real-time notes and transcript ingestion.", "Medium", "May use third-party meeting provider initially"],
                            ["FR-INT-03", "System shall store question order, timestamps, and session event logs.", "High", "Critical for audit and analytics"],
                        ],
                    },
                    {"type": "subheading", "text": "8.5 AI Processing and Insights (FR-AI)"},
                    {
                        "type": "table",
                        "headers": ["ID", "Requirement", "Priority", "Architect Note"],
                        "rows": [
                            ["FR-AI-01", "System shall generate transcript summaries, strengths, risks, and follow-up questions.", "High", "Responses should include model metadata"],
                            ["FR-AI-02", "System shall support confidence markers or evidence-linked rationale.", "High", "Avoid opaque scoring output"],
                            ["FR-AI-03", "System shall allow AI output review, edit, acceptance, or rejection.", "High", "Human review remains mandatory"],
                        ],
                    },
                    {"type": "subheading", "text": "8.6 Reviewer Workflow and Decisions (FR-REV)"},
                    {
                        "type": "table",
                        "headers": ["ID", "Requirement", "Priority", "Architect Note"],
                        "rows": [
                            ["FR-REV-01", "System shall provide structured scorecards for recruiter and manager review.", "High", "Rubric must map to template version used"],
                            ["FR-REV-02", "System shall track manual overrides to AI suggestions.", "High", "Overrides require reason capture"],
                            ["FR-REV-03", "System shall log final disposition and ATS handoff status.", "High", "Decision state machine should be explicit"],
                        ],
                    },
                    {"type": "subheading", "text": "8.7 Reporting and Administration (FR-ADM)"},
                    {
                        "type": "table",
                        "headers": ["ID", "Requirement", "Priority", "Architect Note"],
                        "rows": [
                            ["FR-ADM-01", "System shall provide tenant-level configuration for roles, branding, consent text, and retention policy.", "High", "Avoid hard-coded policy rules"],
                            ["FR-ADM-02", "System shall expose dashboards for funnel conversion, reviewer productivity, and candidate drop-off.", "Medium", "Aggregate data through reporting models, not OLTP joins alone"],
                            ["FR-ADM-03", "System shall expose audit logs for critical actions and exports.", "High", "Logs should be queryable and exportable"],
                        ],
                    },
                ],
            },
            {
                "heading": "9. UI AND UX REQUIREMENTS",
                "items": [
                    {"type": "subheading", "text": "9.1 Design Philosophy"},
                    {
                        "type": "bullets",
                        "items": [
                            "Candidate flows must be low-friction, calm, and mobile-first.",
                            "Recruiter flows must optimize for queue management, status clarity, and rapid comparison.",
                            "AI outputs must be presented as suggestions with evidence, not as deterministic truth.",
                            "All major actions must provide explicit success, warning, or failure feedback.",
                        ],
                    },
                    {"type": "subheading", "text": "9.2 Key Candidate Screens"},
                    {
                        "type": "table",
                        "headers": ["Screen", "Purpose", "Critical UI Elements"],
                        "rows": [
                            ["Invitation Landing", "Set context and expectations", "Job title, duration, consent notice, start button"],
                            ["Candidate Authentication", "Secure session entry", "OTP or secure link validation, expiry handling"],
                            ["Consent and Setup", "Capture legal acknowledgement", "Consent checkbox, device test, retry guidance"],
                            ["Interview Session", "Capture responses", "Question panel, timer, recording indicator, next action"],
                            ["Completion Screen", "Close the flow", "Confirmation, next steps, support path"],
                        ],
                    },
                    {"type": "subheading", "text": "9.3 Key Recruiter and Manager Screens"},
                    {
                        "type": "table",
                        "headers": ["Screen", "Purpose", "Critical UI Elements"],
                        "rows": [
                            ["Dashboard", "View queue and SLA status", "Status cards, filters, alert counts, shortcuts"],
                            ["Template Builder", "Configure governed interviews", "Competency matrix, questions, weights, preview"],
                            ["Candidate Detail", "Review full interview evidence", "Transcript, summary, scorecard, media links"],
                            ["Decision Workspace", "Approve or reject candidate", "Disposition actions, override reason, ATS sync status"],
                            ["Admin Console", "Manage tenant settings", "Roles, policy text, retention windows, integration health"],
                        ],
                    },
                    {"type": "subheading", "text": "9.4 Explicit Anti-Patterns"},
                    {
                        "type": "bullets",
                        "items": [
                            "Do not hide recording state or consent state from candidates.",
                            "Do not bury AI confidence or evidence links behind multiple clicks.",
                            "Do not mix admin controls into day-to-day recruiter screens.",
                            "Do not require desktop-only interactions for core candidate workflows.",
                        ],
                    },
                ],
            },
            {
                "heading": "10. USER ROLES AND PERMISSIONS MATRIX",
                "items": [
                    {
                        "type": "table",
                        "headers": ["Role", "Create Templates", "Run Interviews", "Review Outcomes", "Administer Settings", "Export Data"],
                        "rows": [
                            ["Candidate", "No", "Yes", "No", "No", "No"],
                            ["Recruiter", "Yes", "Indirect", "Yes", "Limited", "Limited"],
                            ["Hiring Manager", "No", "No", "Yes", "No", "Limited"],
                            ["Administrator", "Yes", "No", "Yes", "Yes", "Yes"],
                            ["Compliance Auditor", "No", "No", "Read Only", "Policy Read Only", "Approved Audit Only"],
                        ],
                    },
                    {
                        "type": "paragraph",
                        "text": "Critical actions such as policy changes, template publication, disposition override, and audit export must be individually permissioned and logged.",
                    },
                ],
            },
            {
                "heading": "11. BUSINESS RULES",
                "items": [
                    {
                        "type": "bullets",
                        "items": [
                            "AI recommendations are advisory and cannot finalize a hiring decision without human action.",
                            "Each interview session must be associated with a single template version and job context.",
                            "Any manual override of AI output must store actor, timestamp, before value, after value, and reason.",
                            "Candidate recording must be blocked when required consent is absent or expired.",
                            "Final candidate disposition must include a reason code and reviewer identity.",
                            "Retention and deletion must be enforced per tenant policy and jurisdiction.",
                        ],
                    },
                ],
            },
            {
                "heading": "12. DATA REQUIREMENTS",
                "items": [
                    {"type": "subheading", "text": "12.1 Key Data Entities"},
                    {
                        "type": "table",
                        "headers": ["Entity", "Purpose", "Architect Consideration"],
                        "rows": [
                            ["Tenant", "Customer isolation boundary", "Use tenant-scoped authorization everywhere"],
                            ["User", "Recruiter, manager, admin identity", "Role binding should be explicit and revocable"],
                            ["Candidate", "Subject of interview journey", "PII should be minimized in logs"],
                            ["Interview Template", "Reusable interview configuration", "Version and publish state required"],
                            ["Interview Session", "Execution record of a single interview", "State machine should be explicit"],
                            ["Transcript", "Text output of interview responses", "Store lineage to media and vendor run"],
                            ["Scorecard", "Structured reviewer assessment", "Must map to template rubric schema"],
                            ["Audit Event", "Traceability record", "Immutable append-style model preferred"],
                        ],
                    },
                    {"type": "subheading", "text": "12.2 Data Migration and Import"},
                    {
                        "type": "bullets",
                        "items": [
                            "Import candidate and requisition metadata from ATS.",
                            "Seed interview templates and question banks for pilot roles.",
                            "Backfill historical outcomes only if governance and retention rules permit.",
                            "Use deduplication and identity matching rules during import.",
                        ],
                    },
                ],
            },
            {
                "heading": "13. INTEGRATION REQUIREMENTS",
                "items": [
                    {
                        "type": "table",
                        "headers": ["Integration", "Purpose", "Direction", "Notes"],
                        "rows": [
                            ["ATS", "Candidate and requisition sync", "Bi-directional", "Minimum viable status update and metadata read"],
                            ["Calendar", "Live interview scheduling", "Bi-directional", "Reschedule and cancellation support required"],
                            ["Identity Provider", "SSO and role mapping", "Inbound", "Use SAML or OIDC"],
                            ["Email and SMS", "Notifications and reminders", "Outbound", "Track delivery status and retries"],
                            ["Speech-to-Text", "Transcription generation", "Outbound and callback", "Prefer asynchronous processing"],
                            ["LLM Provider", "Summary and question generation", "Outbound", "Store model version and prompt context"],
                        ],
                    },
                ],
            },
            {
                "heading": "14. COMPLIANCE AND REGULATORY REQUIREMENTS",
                "items": [
                    {
                        "type": "bullets",
                        "items": [
                            "Platform must support configurable consent text and evidence capture per tenant and jurisdiction.",
                            "Platform must support policy-driven retention, purge, and legal hold behavior.",
                            "Exports of interview evidence, candidate data, and audit records must be traceable.",
                            "Access to transcript and media artifacts must be restricted by role, tenant, and business need.",
                        ],
                    },
                ],
            },
            {
                "heading": "15. ARCHITECTURE AND DEVELOPMENT REQUIREMENTS",
                "items": [
                    {"type": "subheading", "text": "15.1 Target Logical Architecture"},
                    {
                        "type": "bullets",
                        "items": [
                            "Presentation layer for candidate, recruiter, manager, and admin experiences",
                            "API or orchestration layer for workflow control and policy enforcement",
                            "Interview service for template, session, and scorecard lifecycle",
                            "AI processing pipeline for transcription, summarization, and question generation",
                            "Notification service for email, SMS, reminder, and retry logic",
                            "Reporting or analytics layer for dashboard aggregation",
                        ],
                    },
                    {"type": "subheading", "text": "15.2 Development Constraints"},
                    {
                        "type": "bullets",
                        "items": [
                            "Core transactional writes must remain deterministic and independent from slow AI provider calls.",
                            "AI processing must be asynchronous and recoverable through queue-based orchestration.",
                            "State transitions for interview sessions and candidate decisions must be explicit and validated server-side.",
                            "All integration points must support timeout, retry, idempotency, and observability.",
                        ],
                    },
                    {"type": "subheading", "text": "15.3 Engineering Environments"},
                    {
                        "type": "table",
                        "headers": ["Environment", "Purpose", "Minimum Requirement"],
                        "rows": [
                            ["Development", "Local feature work", "Mock vendor adapters and sample data"],
                            ["QA", "Integrated functional validation", "Shared test data and stable vendor sandbox"],
                            ["UAT", "Business validation", "Production-like configuration and approval workflow"],
                            ["Production", "Live operations", "Monitoring, audit, backups, and support runbooks"],
                        ],
                    },
                    {"type": "subheading", "text": "15.4 Observability and Supportability"},
                    {
                        "type": "bullets",
                        "items": [
                            "Centralized structured logging for API, workflow, and vendor interaction events",
                            "Metrics for invitation delivery, completion rate, processing latency, and reviewer backlog",
                            "Distributed tracing for critical workflow orchestration where feasible",
                            "Operational dashboards and alert thresholds for failed transcription, high queue delay, and integration outage",
                        ],
                    },
                ],
            },
            {
                "heading": "16. NON-FUNCTIONAL REQUIREMENTS",
                "items": [
                    {
                        "type": "table",
                        "headers": ["Category", "Requirement"],
                        "rows": [
                            ["Security", "SSO, MFA for internal users, encryption in transit and at rest, least privilege, secrets management, and audit logging"],
                            ["Privacy", "Consent capture, retention enforcement, deletion workflows, and data minimization"],
                            ["Performance", "Common recruiter actions should complete within 2 seconds under normal load"],
                            ["Availability", "Target 99.9 percent uptime for core screening workflows"],
                            ["Scalability", "Support burst volume during bulk recruitment without redesign of core services"],
                            ["Accessibility", "Candidate and reviewer interfaces should align to WCAG 2.2 AA baseline"],
                            ["Reliability", "Use retries, dead-letter handling, and idempotent processing for async jobs"],
                            ["Portability", "Vendor adapters should allow provider substitution with limited orchestration change"],
                        ],
                    },
                ],
            },
            {
                "heading": "17. REPORTING AND MIS REQUIREMENTS",
                "items": [
                    {
                        "type": "table",
                        "headers": ["Report", "Audience", "Purpose"],
                        "rows": [
                            ["Interview Funnel Report", "Leadership", "See stage conversion and backlog"],
                            ["Reviewer Productivity", "HR Operations", "Measure turnaround and workload"],
                            ["Candidate Completion Report", "Recruitment Operations", "Identify drop-off by role, region, or device"],
                            ["Compliance Audit Report", "Compliance Team", "Validate consent, access, and export trails"],
                            ["Vendor SLA Report", "Architecture and Operations", "Track latency and failure rates of external services"],
                        ],
                    },
                ],
            },
            {
                "heading": "18. TESTING STRATEGY AND QUALITY REQUIREMENTS",
                "items": [
                    {"type": "subheading", "text": "18.1 Test Levels"},
                    {
                        "type": "bullets",
                        "items": [
                            "Unit tests for business rules, policy checks, and state transitions",
                            "Integration tests for ATS, calendar, notification, and AI vendor adapters",
                            "End-to-end tests for candidate invitation, completion, review, and disposition workflows",
                            "Regression tests for template versioning, consent handling, and audit logging",
                        ],
                    },
                    {"type": "subheading", "text": "18.2 Non-Functional Validation"},
                    {
                        "type": "bullets",
                        "items": [
                            "Performance tests for bulk invitation send and high-volume async processing",
                            "Security tests for authorization boundaries, export restrictions, and secret handling",
                            "Accessibility tests for candidate and reviewer interfaces",
                            "Resilience tests for vendor timeout, retry, and queue recovery scenarios",
                        ],
                    },
                    {"type": "subheading", "text": "18.3 UAT Focus Areas"},
                    {
                        "type": "bullets",
                        "items": [
                            "Candidate experience clarity",
                            "Recruiter dashboard usefulness",
                            "Scorecard correctness and override traceability",
                            "ATS handoff accuracy",
                            "Policy and retention behavior",
                        ],
                    },
                ],
            },
            {
                "heading": "19. KEY WORKFLOW DIAGRAMS (TEXTUAL)",
                "items": [
                    {"type": "subheading", "text": "19.1 Asynchronous Candidate Journey"},
                    {
                        "type": "numbered",
                        "items": [
                            "Recruiter triggers invitation",
                            "Candidate authenticates and accepts consent",
                            "Candidate completes interview questions",
                            "Media is stored and transcription job is queued",
                            "AI summary is generated and linked to evidence",
                            "Recruiter reviews and records disposition",
                            "Decision is synchronized to ATS",
                        ],
                    },
                    {"type": "subheading", "text": "19.2 Live Interview Path"},
                    {
                        "type": "numbered",
                        "items": [
                            "Recruiter schedules live session",
                            "Calendar invite and reminders are sent",
                            "Candidate joins session",
                            "Notes and transcript data are ingested",
                            "Reviewer completes scorecard",
                            "Final summary and decision are stored",
                        ],
                    },
                ],
            },
            {
                "heading": "20. RISKS AND MITIGATIONS",
                "items": [
                    {
                        "type": "table",
                        "headers": ["Risk", "Impact", "Mitigation"],
                        "rows": [
                            ["AI bias or unreliable recommendation", "High", "Keep human review mandatory and require evidence-linked outputs"],
                            ["Vendor outage or high latency", "High", "Use async queues, retries, and provider abstraction"],
                            ["Poor candidate user experience", "Medium", "Design mobile-first flows with clear guidance and recovery states"],
                            ["Cross-tenant data leakage", "High", "Enforce tenant isolation at auth, query, and storage layers"],
                            ["Compliance breach due to retention or consent gaps", "High", "Make policy configurable and audit every critical action"],
                        ],
                    },
                ],
            },
            {
                "heading": "21. OPEN ISSUES AND DECISIONS PENDING",
                "items": [
                    {
                        "type": "bullets",
                        "items": [
                            "Confirm MVP scope for live interview support versus asynchronous-only launch.",
                            "Confirm initial ATS and calendar vendors in rollout phase one.",
                            "Confirm approved AI vendors and fallback policy.",
                            "Confirm legal retention windows by operating region.",
                            "Confirm whether multilingual interview support is MVP or later phase.",
                        ],
                    },
                ],
            },
            {
                "heading": "22. APPENDIX",
                "items": [
                    {"type": "subheading", "text": "22.1 Glossary"},
                    {
                        "type": "bullets",
                        "items": [
                            "ATS: Applicant Tracking System",
                            "LLM: Large Language Model",
                            "STT: Speech-to-text",
                            "Scorecard: Structured reviewer assessment sheet",
                            "Tenant: A logically isolated customer environment within the platform",
                        ],
                    },
                    {"type": "subheading", "text": "22.2 Source Reference"},
                    {
                        "type": "paragraph",
                        "text": "This BRD was structured using the enterprise document style of the uploaded ASG-Vasan reference BRD while adapting the content to the AI Interview Platform domain.",
                    },
                ],
            },
        ],
    }


def configure_docx_styles(document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.bold = True
        style.font.size = Pt(size)


def add_docx_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell_fill(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    document.add_paragraph()


def build_docx(model):
    document = Document()
    configure_docx_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{model['title']}\n{model['subtitle']}")
    title_run.bold = True
    title_run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(model["tagline"]).italic = True

    add_docx_table(document, ["Field", "Value"], model["meta"])

    for section in model["sections"]:
        document.add_heading(section["heading"], level=1)
        for item in section["items"]:
            item_type = item["type"]
            if item_type == "subheading":
                document.add_heading(item["text"], level=2)
            elif item_type == "paragraph":
                document.add_paragraph(item["text"])
            elif item_type == "bullets":
                for bullet in item["items"]:
                    document.add_paragraph(bullet, style="List Bullet")
            elif item_type == "numbered":
                for step in item["items"]:
                    document.add_paragraph(step, style="List Number")
            elif item_type == "table":
                add_docx_table(document, item["headers"], item["rows"])

    document.save(DOCX_PATH)


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=22, leading=26, spaceAfter=10))
    styles.add(ParagraphStyle(name="SubCenter", parent=styles["Normal"], alignment=TA_CENTER, fontName="Helvetica-Oblique", fontSize=10, leading=13, textColor=colors.HexColor("#475569"), spaceAfter=12))
    styles.add(ParagraphStyle(name="Section", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=17, textColor=colors.HexColor("#0f172a"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="Subsection", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=13, textColor=colors.HexColor("#1e293b"), spaceBefore=5, spaceAfter=3))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=11.5, spaceAfter=4))
    styles.add(ParagraphStyle(name="SmallBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.1, leading=9.5, spaceAfter=3))
    return styles


def pdf_table(styles, headers, rows, widths):
    data = [[Paragraph(str(header), styles["SmallBody"]) for header in headers]]
    for row in rows:
        data.append([Paragraph(str(value), styles["SmallBody"]) for value in row])
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def get_col_widths(column_count):
    total_width = 7.0 * inch
    if column_count == 2:
        return [1.9 * inch, 5.1 * inch]
    if column_count == 3:
        return [1.5 * inch, 1.1 * inch, 4.4 * inch]
    if column_count == 4:
        return [1.2 * inch, 1.6 * inch, 1.0 * inch, 3.2 * inch]
    if column_count == 5:
        return [1.1 * inch, 1.2 * inch, 1.0 * inch, 2.7 * inch, 1.0 * inch]
    if column_count == 6:
        return [1.15 * inch, 0.9 * inch, 0.9 * inch, 1.0 * inch, 1.0 * inch, 1.05 * inch]
    return [total_width / column_count] * column_count


def build_pdf(model):
    styles = build_pdf_styles()
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=0.58 * inch,
        rightMargin=0.58 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    story = []

    story.append(Paragraph(model["title"], styles["TitleCenter"]))
    story.append(Paragraph(model["subtitle"], styles["TitleCenter"]))
    story.append(Paragraph(model["tagline"], styles["SubCenter"]))
    story.append(pdf_table(styles, ["Field", "Value"], model["meta"], get_col_widths(2)))
    story.append(Spacer(1, 0.12 * inch))

    for section in model["sections"]:
        story.append(Paragraph(section["heading"], styles["Section"]))
        for item in section["items"]:
            item_type = item["type"]
            if item_type == "subheading":
                story.append(Paragraph(item["text"], styles["Subsection"]))
            elif item_type == "paragraph":
                story.append(Paragraph(item["text"], styles["Body"]))
            elif item_type == "bullets":
                for bullet in item["items"]:
                    story.append(Paragraph(f"- {bullet}", styles["Body"]))
            elif item_type == "numbered":
                for index, step in enumerate(item["items"], start=1):
                    story.append(Paragraph(f"{index}. {step}", styles["Body"]))
            elif item_type == "table":
                story.append(pdf_table(styles, item["headers"], item["rows"], get_col_widths(len(item["headers"]))))
                story.append(Spacer(1, 0.08 * inch))

    def add_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#475569"))
        canvas.drawRightString(A4[0] - 0.58 * inch, 0.38 * inch, f"Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=add_footer, onLaterPages=add_footer)


if __name__ == "__main__":
    model = get_document_model()
    build_docx(model)
    build_pdf(model)
    print(f"Created {DOCX_PATH}")
    print(f"Created {PDF_PATH}")
