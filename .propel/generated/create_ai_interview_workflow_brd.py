from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "AI_Interview_Workflow_BRD.docx"
PDF_PATH = ROOT / "AI_Interview_Workflow_BRD.pdf"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def setup_doc(document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.bold = True
        style.font.size = Pt(size)


def add_table_docx(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    document.add_paragraph()


def add_bullets_docx(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered_docx(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def workflow_steps():
    return [
        [
            "WF-01",
            "Applicant Registration",
            "Applicant",
            "New applicant opens hiring portal",
            "Capture basic profile, verify email or phone with OTP, persist consent",
            "Verified candidate account",
            "If OTP fails 3 times, account lock for 15 minutes",
            "P95 registration completion < 90 seconds",
        ],
        [
            "WF-02",
            "Login",
            "Applicant, Recruiter, HR, Interviewer",
            "Registered user submits credentials",
            "Authenticate, authorize by role, issue session token",
            "Role-specific dashboard",
            "Invalid credentials increment lock counter and return generic error",
            "P95 login response < 1.5 seconds",
        ],
        [
            "WF-03",
            "Apply for Job",
            "Applicant",
            "Candidate selects open requisition",
            "Validate eligibility and duplicate application rules",
            "Draft or submitted application",
            "If duplicate within cooling period, block with reason code",
            "P95 search and apply action < 2 seconds",
        ],
        [
            "WF-04",
            "Upload Resume",
            "Applicant",
            "Application in draft or submitted state",
            "Accept PDF or DOCX, malware scan, parse document",
            "Resume stored and parsed profile",
            "On parser failure, create retry job and alert candidate",
            "Upload success rate > 99.5 percent",
        ],
        [
            "WF-05",
            "AI Resume Screening",
            "System",
            "Resume parsing complete",
            "Run skill extraction, match scoring, confidence and rationale generation",
            "Screening score and recommendation",
            "If model unavailable, queue retry then fallback to recruiter-only screening",
            "P95 screening completion < 4 minutes",
        ],
        [
            "WF-06",
            "HR Review",
            "HR Reviewer",
            "AI result available",
            "Review AI rationale, resume highlights, requisition fit, and conflict flags",
            "Shortlist or reject decision",
            "If reviewer skips mandatory fields, prevent submission",
            "Review SLA < 24 business hours",
        ],
        [
            "WF-07",
            "Shortlisted Decision Branch",
            "HR Reviewer",
            "Review decision submitted",
            "Route candidate to rejection or interview flow",
            "Workflow status updated",
            "If routing service fails, preserve decision and retry branch routing",
            "State transition success > 99.9 percent",
        ],
        [
            "WF-08",
            "Rejection Communication",
            "System",
            "Not shortlisted",
            "Render rejection template, send email, track delivery",
            "Communication log entry",
            "If email bounces, create recruiter follow-up task",
            "Mail delivery confirmation within 5 minutes",
        ],
        [
            "WF-09",
            "Interview Process Initiation",
            "Recruiter",
            "Shortlisted candidate",
            "Collect availability windows and move candidate to interview lifecycle",
            "Interview plan shell",
            "If calendar sync fails, fallback to manual slot assignment",
            "Plan creation < 10 minutes",
        ],
        [
            "WF-10",
            "Experience Check",
            "Recruiter",
            "Interview process initiated",
            "Classify candidate as fresher or experienced using policy",
            "Path assignment",
            "If data missing, hold candidate in pending verification queue",
            "Classification error < 1 percent",
        ],
        [
            "WF-11",
            "Fresher Path Aptitude Test",
            "Applicant",
            "Path assigned as fresher",
            "Launch aptitude assessment with proctoring and timer",
            "Aptitude score",
            "If disconnected, resume once within policy window",
            "Assessment completion > 95 percent",
        ],
        [
            "WF-12",
            "Experienced Path Technical Interview",
            "Technical Interviewer",
            "Path assigned as experienced",
            "Conduct structured interview and submit rubric-based scorecard",
            "Technical recommendation",
            "If scorecard incomplete, block sign-off and notify interviewer",
            "Feedback submitted within 12 hours",
        ],
        [
            "WF-13",
            "Programming Assessment",
            "Applicant",
            "Aptitude or initial technical gate completed",
            "Run coding challenge, auto-evaluate tests, collect plagiarism signals",
            "Programming score",
            "If judge service fails, switch to manual evaluation queue",
            "Auto evaluation success > 99 percent",
        ],
        [
            "WF-14",
            "Technical Interview",
            "Technical Interviewer",
            "Programming score available",
            "Assess design, problem solving, and communication",
            "Technical panel recommendation",
            "On no-show, auto-reschedule once then mark as no-show",
            "Interview closure within 2 business days",
        ],
        [
            "WF-15",
            "HR Round",
            "HR Manager",
            "Technical recommendation available",
            "Evaluate culture fit, compensation fit, and policy checks",
            "HR recommendation",
            "If compensation band mismatch, route to compensation exception flow",
            "HR round closure within 1 business day",
        ],
        [
            "WF-16",
            "Final Decision",
            "HR Manager",
            "All mandatory stages complete",
            "Approve hire or reject with reason category",
            "Final decision record",
            "If approvals incomplete, keep status at pending-approval",
            "Decision latency < 8 business hours",
        ],
        [
            "WF-17",
            "Offer or Final Rejection",
            "System and Recruiter",
            "Final decision recorded",
            "Generate offer packet or final rejection communication",
            "Candidate notified and workflow closed",
            "If candidate does not respond by deadline, auto-trigger reminder",
            "Offer response tracking daily",
        ],
    ]


def functional_modules():
    return [
        (
            "FR-REG: Registration, Login, and Candidate Identity",
            [
                ["FR-REG-01", "System shall support candidate self-registration with email or phone OTP verification.", "High", "Identity boundary starts here; reject unverifiable identities.", "User receives verified status and can log in."],
                ["FR-REG-02", "System shall enforce configurable password policy and lockout thresholds.", "High", "Control brute-force risk through policy parameters.", "After 5 failures, account lock event is logged."],
                ["FR-REG-03", "System shall capture and version consent for privacy and AI processing notices.", "High", "Consent must be immutable and time-stamped.", "Consent version appears in candidate audit timeline."],
                ["FR-REG-04", "System shall support role-based login for applicant, recruiter, interviewer, HR manager, and admin.", "High", "Use role claims, not UI-only role checks.", "Users only see permitted navigation and APIs."],
                ["FR-REG-05", "System shall support optional SSO for internal users.", "Medium", "Integrate with enterprise IdP for staff roles.", "SSO users can authenticate without local password."],
            ],
        ),
        (
            "FR-APP: Job Search, Application, and Resume Capture",
            [
                ["FR-APP-01", "System shall display requisitions with eligibility criteria, status, and location.", "High", "Single source from ATS or requisition service.", "Candidate can filter and apply to open jobs."],
                ["FR-APP-02", "System shall prevent duplicate applications for the same requisition within a configurable cooling period.", "High", "Avoid duplicate funnel noise and recruiter effort.", "Duplicate attempts show reason and cooling period end date."],
                ["FR-APP-03", "System shall accept resume uploads in PDF and DOCX with malware scanning.", "High", "Reject unsupported format before parse call.", "Upload returns accepted and scanning-complete status."],
                ["FR-APP-04", "System shall parse resumes and map extracted fields to candidate profile attributes.", "High", "Parser output must preserve source confidence.", "At least name, email, experience, skills, education are mapped."],
                ["FR-APP-05", "System shall allow candidate edits to extracted profile before final submission.", "Medium", "Human correction improves downstream matching quality.", "Manual edits are audit-tracked separately from parser output."],
            ],
        ),
        (
            "FR-AI: AI Screening and Explainability",
            [
                ["FR-AI-01", "System shall compute match score using required skills, experience, and role fit.", "High", "Keep scoring schema versioned for reproducibility.", "Score and score-version are stored with timestamp."],
                ["FR-AI-02", "System shall store top contributing factors and gaps for HR explainability view.", "High", "No black-box decision without rationale context.", "Reviewer sees positives, gaps, and confidence."],
                ["FR-AI-03", "System shall flag low-confidence results for mandatory human review.", "High", "Prevent low-certainty automation from driving rejection.", "Low confidence routes candidate to manual queue."],
                ["FR-AI-04", "System shall support configurable score thresholds per job family.", "High", "Threshold policy owned by HR operations.", "Threshold change effective date and author are logged."],
                ["FR-AI-05", "System shall support model fallback mode when AI service is degraded.", "Medium", "Business continuity with recruiter-led review.", "System enters fallback state and notifies operations."],
            ],
        ),
        (
            "FR-REV: HR Review, Shortlist, and Routing",
            [
                ["FR-REV-01", "System shall provide HR review dashboard with candidate queue, filters, and aging indicators.", "High", "Aging ensures SLA compliance visibility.", "Queue can be filtered by requisition, stage, and SLA breach."],
                ["FR-REV-02", "System shall require shortlist or reject decision with mandatory reason codes.", "High", "Reason taxonomy enables analytics.", "Decision submission blocked without required fields."],
                ["FR-REV-03", "System shall trigger rejection communication automatically for not shortlisted candidates.", "High", "Ensure timely closure and candidate transparency.", "Communication event appears in timeline and audit logs."],
                ["FR-REV-04", "System shall classify shortlisted candidates as fresher or experienced by policy.", "High", "Routing policy must be deterministic.", "Classified path visible and editable by authorized role."],
                ["FR-REV-05", "System shall support recruiter override with justification and approval for exceptional routing.", "Medium", "Controlled override for edge cases.", "Override request and approver are captured."],
            ],
        ),
        (
            "FR-INT: Assessments and Interviews",
            [
                ["FR-INT-01", "System shall schedule aptitude, coding, and interview stages with timezone-aware slots.", "High", "Store UTC and local timezone metadata.", "Candidate and panel receive correctly localized schedules."],
                ["FR-INT-02", "System shall enforce fresher path with mandatory aptitude stage before coding.", "High", "Branch integrity rule.", "Fresher cannot open coding link before aptitude completion."],
                ["FR-INT-03", "System shall allow experienced path to begin with technical interview per policy.", "High", "Configurable workflow while preserving governance.", "Experienced candidate can proceed as configured."],
                ["FR-INT-04", "System shall ingest aptitude and coding scores from external assessment provider.", "High", "Use signed callbacks and idempotent ingestion.", "Duplicate callback does not create duplicate score rows."],
                ["FR-INT-05", "System shall capture interviewer scorecards with rubric sections and recommendation.", "High", "Structured scorecards improve consistency.", "Scorecard includes mandatory rubric dimensions and comments."],
                ["FR-INT-06", "System shall support no-show, reschedule, and cancellation states with reason codes.", "Medium", "Operational realism and metrics quality.", "State changes are tracked and visible in timeline."],
            ],
        ),
        (
            "FR-DEC: Final Decision, Offer Governance, and Closure",
            [
                ["FR-DEC-01", "System shall permit final decision only when mandatory stages are complete.", "High", "Gate control prevents premature closure.", "Decision button disabled until prerequisites are met."],
                ["FR-DEC-02", "System shall support decision outcomes: offer, reject, hold, and withdraw.", "High", "Comprehensive closure states for operations.", "Outcome reflected in candidate and requisition metrics."],
                ["FR-DEC-03", "System shall enforce approval matrix for offer release based on compensation band.", "High", "Risk control for compensation governance.", "Offer cannot be sent without required approvers."],
                ["FR-DEC-04", "System shall produce auditable decision summary with stage scores and rationale.", "High", "Essential for compliance and dispute handling.", "Summary export available in recruiter console."],
                ["FR-DEC-05", "System shall support candidate response tracking and auto-reminders for pending offers.", "Medium", "Improve offer-to-join conversion.", "Reminder jobs run until response or expiry."],
            ],
        ),
        (
            "FR-COM: Communication and Notifications",
            [
                ["FR-COM-01", "System shall provide tokenized email templates for offer, rejection, interview invite, and reminders.", "High", "Template governance should be admin-controlled.", "Template preview renders correctly with sample data."],
                ["FR-COM-02", "System shall log all outbound communication with provider status.", "High", "Communication traceability is mandatory.", "Delivery, bounce, and retry statuses are visible."],
                ["FR-COM-03", "System shall retry transient notification failures with exponential backoff.", "High", "Protect against temporary provider instability.", "Retries stop after max attempts and create task."],
                ["FR-COM-04", "System shall support multi-channel alerts for internal users (email and in-app).", "Medium", "Reduce missed internal actions.", "Recruiter receives both in-app alert and email for critical SLA events."],
                ["FR-COM-05", "System shall support localization-ready template content.", "Low", "Prepare for multi-region rollout.", "Template model includes locale and fallback locale."],
            ],
        ),
    ]


def build_docx():
    document = Document()
    setup_doc(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI Interview Application\nBusiness Requirements Document")
    title_run.bold = True
    title_run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Architect perspective with implementation-ready detail and workflow governance").italic = True

    add_table_docx(
        document,
        ["Field", "Value"],
        [
            ["Version", "2.0"],
            ["Prepared Date", date.today().isoformat()],
            ["Document Type", "Workflow-based enterprise BRD"],
            ["Source Workflow", "Registration to Offer/Rejection with fresher and experienced branches"],
            ["Status", "Draft for architecture, product, and operations sign-off"],
            ["Organization Name", "<Your Company Name>"],
            ["Product Name", "<Your AI Interview Product Name>"],
        ],
    )

    document.add_heading("0. Organization, Branding, and Ownership", level=1)
    document.add_paragraph(
        "This section formalizes ownership, governance boundaries, and operating context for the AI Interview Application. "
        "The platform is positioned as a hiring execution system that combines AI-assisted screening with human governance checkpoints, "
        "ensuring speed, quality, fairness, and auditability."
    )
    add_table_docx(
        document,
        ["Attribute", "Value"],
        [
            ["Company Legal Name", "<Your Company Legal Name>"],
            ["Product Name", "<Your Product Name>"],
            ["Primary Business Unit", "Talent Acquisition and HR Technology"],
            ["Executive Sponsor", "<Chief Human Resources Officer>"],
            ["Business Owner", "Head of Talent Acquisition"],
            ["Technical Owner", "Solution Architect"],
            ["Operations Owner", "HR Operations Manager"],
            ["Security Owner", "Information Security Lead"],
            ["Support Model", "L1 HR Ops, L2 Product Support, L3 Engineering"],
        ],
    )

    document.add_heading("0.1 Project Charter Context", level=2)
    add_bullets_docx(
        document,
        [
            "Problem: Manual recruitment funnels are inconsistent, slow, and weakly auditable across high-volume hiring programs.",
            "Vision: Build a governed AI-assisted hiring platform that improves decision quality while reducing time-to-hire.",
            "Target Outcome: Implement measurable process reliability across screening, interviews, final decisions, and candidate communication.",
            "Delivery Principle: Human-in-the-loop decisions for all adverse outcomes and compensation approvals.",
            "Success Horizon: Stable production rollout with SLA compliance and controlled model governance.",
        ],
    )

    document.add_heading("0.2 Delivery Timeline Baseline", level=2)
    add_table_docx(
        document,
        ["Phase", "Duration", "Primary Deliverables", "Exit Criteria"],
        [
            ["Phase 1: Discovery and Design", "4 weeks", "BRD sign-off, architecture blueprint, data model baseline", "All critical decisions and open issues triaged"],
            ["Phase 2: Core Build", "8 weeks", "Registration, application, screening, review, routing modules", "Core workflow E2E in SIT"],
            ["Phase 3: Interview and Decision", "6 weeks", "Assessments, interviews, final decision governance, communication", "UAT pass with business sign-off"],
            ["Phase 4: Stabilization and Go-Live", "4 weeks", "Performance hardening, observability, runbooks, training", "Production readiness checklist complete"],
        ],
    )

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "The AI Interview Application addresses a business-critical gap in modern hiring operations: fragmented processes, variable reviewer quality, "
        "and delayed candidate closure. The proposed platform centralizes candidate progression, embeds stage-level governance, and provides explainable AI screening "
        "without removing human accountability."
    )
    document.add_paragraph(
        "From an architect perspective, the platform is designed as a modular workflow system with deterministic state transitions, policy-driven branching, "
        "and robust integration points for ATS, assessment engines, calendar systems, and email providers."
    )
    document.add_paragraph(
        "Business value is measured through reduced time-to-hire, better shortlist precision, lower interview no-show impact, and improved candidate communication SLAs. "
        "Technical value is measured through service reliability, secure data handling, and audit-grade traceability."
    )
    add_table_docx(
        document,
        ["Objective", "Current State", "Target State", "Measurement"],
        [
            ["Reduce hiring cycle time", "Manual stage coordination", "Workflow automation and SLA alerts", "Median time-to-hire"],
            ["Improve shortlist quality", "Subjective screening", "AI score plus HR validation", "Interview-to-offer conversion"],
            ["Increase candidate experience", "Inconsistent updates", "Template-driven stage communication", "Candidate communication SLA"],
            ["Strengthen auditability", "Scattered records", "Centralized immutable event logs", "Decision trace completeness"],
        ],
    )

    document.add_heading("2. Workflow Scope and Detailed Process", level=1)
    document.add_paragraph(
        "The workflow below defines the canonical hiring lifecycle and includes explicit entry/exit conditions, failure paths, and stage-level SLAs. "
        "Any process variation in production must be represented as a controlled policy configuration, not ad-hoc manual workarounds."
    )
    add_table_docx(
        document,
        ["Step ID", "Step Name", "Actor", "Trigger", "System Action", "Output", "Failure Path", "SLA"],
        workflow_steps(),
    )

    document.add_heading("2.1 In Scope", level=2)
    add_bullets_docx(
        document,
        [
            "Candidate account lifecycle, identity verification, and consent capture",
            "Application submission, resume parsing, and AI-assisted screening",
            "HR review and shortlist governance with reason taxonomy",
            "Path branching for fresher and experienced candidates",
            "Aptitude, coding, technical interview, and HR round orchestration",
            "Final decision approval workflow and communication automation",
            "Reporting, audit, security controls, and operations monitoring",
        ],
    )

    document.add_heading("2.2 Out of Scope", level=2)
    add_bullets_docx(
        document,
        [
            "Payroll onboarding and post-offer employee lifecycle",
            "Background verification execution system (only status integration supported)",
            "Internal workforce planning and requisition budgeting modules",
            "Video interview recording storage (unless separately approved by compliance)",
        ],
    )

    document.add_heading("3. Functional Requirements", level=1)
    document.add_paragraph(
        "Functional requirements are grouped by domain module. Each requirement includes priority, architecture guidance, and acceptance criteria suitable "
        "for decomposition into epics, user stories, and test cases."
    )
    for module_name, rows in functional_modules():
        document.add_heading(module_name, level=2)
        add_table_docx(
            document,
            ["ID", "Requirement", "Priority", "Architect Note", "Acceptance Criteria"],
            rows,
        )

    document.add_heading("4. Business Rules and Decision Logic", level=1)
    add_table_docx(
        document,
        ["Rule ID", "Condition (IF)", "Outcome (THEN)", "Override Policy", "Evidence"],
        [
            ["BR-01", "Resume is not uploaded", "Candidate cannot proceed to AI screening", "No override", "Workflow event log"],
            ["BR-02", "AI confidence < configured threshold", "Mandatory HR manual review", "Only HR manager can bypass", "Decision history"],
            ["BR-03", "Candidate rejected at shortlist", "Rejection email must be sent within SLA", "None", "Communication log"],
            ["BR-04", "Candidate marked fresher", "Aptitude must complete before coding assessment", "No override", "Assessment state"],
            ["BR-05", "Candidate marked experienced", "Technical interview can precede coding if policy enabled", "Recruiter override with approval", "Routing policy record"],
            ["BR-06", "Mandatory stage scorecard missing", "Final decision action disabled", "No override", "UI validation log"],
            ["BR-07", "Offer compensation exceeds threshold band", "Additional approver required", "Compensation exception flow", "Approval matrix log"],
            ["BR-08", "Assessment callback duplicated", "System processes once using idempotency key", "No override", "Integration log"],
            ["BR-09", "Interview no-show occurs", "Auto-reschedule once then mark no-show", "Recruiter manual reopen", "Timeline event"],
            ["BR-10", "Email delivery bounces", "Create recruiter task for alternate outreach", "No override", "Bounce status"],
            ["BR-11", "Candidate withdraws", "Workflow closes with withdraw reason", "No override", "Closure record"],
            ["BR-12", "HR reviewer and interviewer conflict flagged", "Decision routed to HR manager", "No override", "Conflict flag audit"],
            ["BR-13", "SLA breach detected", "Escalation alert sent to owner", "SLA suppress for approved exception", "SLA dashboard"],
            ["BR-14", "PII field update requested by internal role", "Allow only authorized role and audit change", "No override", "PII access log"],
            ["BR-15", "Candidate requests data deletion", "Retain legal-minimum records and anonymize operational data", "Compliance sign-off required", "Retention action log"],
            ["BR-16", "Model version changes", "New screenings use new version; prior records remain immutable", "No override", "Model registry"],
            ["BR-17", "Requisition closes", "Pending applications move to closed-not-selected", "HR manager can reopen requisition", "Requisition state history"],
            ["BR-18", "Final decision is reject after interviews", "Mandatory structured rejection reason required", "No override", "Final decision payload"],
        ],
    )

    document.add_heading("5. UI/UX Requirements", level=1)
    document.add_heading("5.1 Experience Principles", level=2)
    add_bullets_docx(
        document,
        [
            "Decision-first dashboards: reviewers should see next required action before secondary details.",
            "Progressive disclosure: show candidate detail depth as user moves from queue to final decision.",
            "Error transparency: each blocking error must include actionable guidance.",
            "Accessibility by design: keyboard navigation, screen-reader labels, sufficient contrast, and focus order.",
            "Status clarity: workflow stage and SLA state are always visible without opening nested views.",
        ],
    )

    document.add_heading("5.2 Screen Catalog", level=2)
    add_table_docx(
        document,
        ["Screen", "Primary Actor", "Mandatory Components", "Error/Empty States", "Telemetry"],
        [
            ["Registration", "Applicant", "Identity fields, OTP flow, consent checkboxes", "OTP failure, duplicate account, timeout", "Registration start/completion"],
            ["Login", "All roles", "Credential inputs, SSO button, role landing", "Invalid credentials, lockout", "Login success/failure"],
            ["Job Search and Apply", "Applicant", "Filters, role card, eligibility panel, apply CTA", "No open roles, duplicate apply block", "Job view-to-apply conversion"],
            ["Resume Upload", "Applicant", "Drag-drop upload, file validation, parse status", "Unsupported file, parser failure", "Upload and parse duration"],
            ["HR Review Dashboard", "HR Reviewer", "Queue table, AI score, confidence, rationale panel", "No pending items, SLA breach marker", "Review throughput and aging"],
            ["Assessment Console", "Applicant", "Timer, question navigation, autosave", "Network drop, timer expiry", "Drop-off and completion"],
            ["Interview Scorecard", "Interviewer", "Rubric sections, notes, recommendation", "Incomplete rubric block", "Scorecard completion time"],
            ["Decision Workbench", "HR Manager", "Stage summary, approval chain, final action", "Missing mandatory stages", "Decision latency"],
            ["Communication Center", "Recruiter", "Template preview, send history, status", "Provider outage fallback", "Delivery and bounce rates"],
        ],
    )

    document.add_heading("5.3 Device and Browser Support", level=2)
    add_table_docx(
        document,
        ["Client Type", "Supported", "Minimum Resolution", "Notes"],
        [
            ["Desktop Browser", "Yes", "1366x768", "Primary recruiter and HR interface"],
            ["Laptop Browser", "Yes", "1280x720", "Responsive layout mandatory"],
            ["Tablet Browser", "Partial", "1024x768", "Reviewer dashboard read-first mode"],
            ["Mobile Browser", "Candidate-first", "390x844", "Candidate apply and test reminders"],
        ],
    )

    document.add_heading("5.4 UI Anti-Patterns to Avoid", level=2)
    add_bullets_docx(
        document,
        [
            "Hidden workflow transitions that change candidate status without explicit confirmation",
            "Overloaded dashboard cards with ambiguous action ownership",
            "Color-only status communication without icon or text fallback",
            "Long forms without autosave for candidate-facing stages",
            "Modal chains that interrupt scorecard completion",
        ],
    )

    document.add_heading("6. Architecture and Development Requirements", level=1)
    document.add_paragraph(
        "Architecture follows layered design with bounded services and event-driven integration for long-running tasks. "
        "Critical state transitions are synchronous and strongly validated; non-blocking tasks are event-queued with retries and dead-letter capture."
    )

    document.add_heading("6.0 Technology Stack", level=2)
    add_table_docx(
        document,
        ["Layer", "Technology", "Purpose", "Notes"],
        [
            ["Frontend", "React", "Candidate and internal hiring user interfaces", "Component-driven UI with responsive screens"],
            ["Backend", "Node.js", "Core APIs, workflow orchestration, and integration handlers", "Primary application runtime for business services"],
            ["AI Integration", "Python", "Resume parsing, AI screening pipelines, and model-serving adapters", "Exposed via service APIs or async workers"],
            ["Database", "PostgreSQL", "Transactional storage for candidates, applications, stages, and decisions", "Relational model with audit-friendly schema"],
        ],
    )

    document.add_heading("6.1 Component View", level=2)
    add_table_docx(
        document,
        ["Component", "Responsibility", "Input", "Output", "Technology Constraint"],
        [
            ["Identity Service", "Authentication and authorization", "Credentials or SSO token", "Session token and role claims", "Must support RBAC and MFA hooks"],
            ["Candidate Service", "Profile and application lifecycle", "Applicant actions", "Candidate and application states", "PII encryption required"],
            ["Resume and AI Service", "Resume parse and match scoring", "Resume document and requisition", "Score, confidence, rationale", "Model versioning mandatory"],
            ["Workflow Orchestrator", "Stage transitions and guards", "Events and API commands", "Deterministic status changes", "No hidden transitions"],
            ["Assessment Adapter", "External test platform integration", "Launch requests and callbacks", "Scores and completion states", "Idempotent callback processing"],
            ["Interview Service", "Scheduling and scorecards", "Calendar events and interviewer input", "Interview outcomes", "Timezone-safe scheduling"],
            ["Decision Service", "Final approvals and closure", "Stage outputs", "Offer/reject/hold outcome", "Approval matrix enforcement"],
            ["Notification Service", "Template rendering and delivery", "Business events", "Email and in-app notifications", "Retry with backoff"],
            ["Audit and Reporting Service", "Immutable event records and KPIs", "All critical actions", "Audit trail and reports", "Append-only logs"],
        ],
    )

    document.add_heading("6.2 API Surface (Representative)", level=2)
    add_table_docx(
        document,
        ["API", "Method", "Purpose", "Authorization", "Failure Contract"],
        [
            ["/api/v1/auth/login", "POST", "Authenticate user", "Public", "401 generic auth error, lockout metadata"],
            ["/api/v1/candidates/{id}/applications", "POST", "Submit application", "Applicant", "409 duplicate application rule"],
            ["/api/v1/resumes/upload", "POST", "Upload and scan resume", "Applicant", "415 invalid format, 422 parsing blocked"],
            ["/api/v1/screening/run/{applicationId}", "POST", "Trigger screening", "System or recruiter", "503 model degraded fallback path"],
            ["/api/v1/reviews/{applicationId}/decision", "POST", "Shortlist or reject", "HR reviewer", "400 missing mandatory reason"],
            ["/api/v1/assessments/callback", "POST", "Ingest external result", "Signed callback", "409 duplicate idempotency key"],
            ["/api/v1/decisions/{candidateId}", "POST", "Publish final decision", "HR manager", "412 missing prerequisite stage"],
            ["/api/v1/communications/send", "POST", "Trigger communication", "System and recruiter", "502 provider unavailable with retry id"],
        ],
    )

    document.add_heading("6.3 Event and Queue Design", level=2)
    add_table_docx(
        document,
        ["Event", "Producer", "Consumer", "Retry Policy", "Dead Letter Handling"],
        [
            ["ResumeUploaded", "Candidate Service", "Resume and AI Service", "3 retries, exponential backoff", "Route to manual parse queue"],
            ["ScreeningCompleted", "Resume and AI Service", "HR Review Service", "2 retries", "Ops alert and replay action"],
            ["AssessmentCompleted", "Assessment Adapter", "Workflow Orchestrator", "Idempotent with 5 retries", "Store payload for manual replay"],
            ["DecisionPublished", "Decision Service", "Notification and Reporting", "3 retries", "Escalate to support queue"],
            ["CommunicationFailed", "Notification Service", "Recruiter Task Service", "1 retry", "Manual intervention task"],
        ],
    )

    document.add_heading("6.4 Deployment and Security Zones", level=2)
    add_table_docx(
        document,
        ["Zone", "Contents", "Inbound Controls", "Outbound Controls"],
        [
            ["Public Zone", "Candidate web app", "WAF, rate limits, bot mitigation", "Only to API gateway"],
            ["Application Zone", "Core services and orchestrator", "Mutual TLS from gateway", "Allowlisted calls to integrations"],
            ["Data Zone", "Primary DB, audit store, object storage", "Private subnet only", "No direct public egress"],
            ["Operations Zone", "Monitoring, alerting, admin console", "SSO and privileged access", "Restricted diagnostics endpoints"],
        ],
    )

    document.add_heading("7. Data and Integration Requirements", level=1)
    document.add_heading("7.1 Core Entities", level=2)
    add_table_docx(
        document,
        ["Entity", "Primary Key", "Critical Attributes", "Retention", "Notes"],
        [
            ["Candidate", "candidate_id", "name, email, phone, consent_version", "24 months", "PII protected and encrypted"],
            ["Requisition", "requisition_id", "title, department, mandatory_skills, status", "As per ATS policy", "Source of truth can be ATS"],
            ["Application", "application_id", "candidate_id, requisition_id, stage, stage_timestamp", "24 months", "One candidate can have multiple applications"],
            ["Resume", "resume_id", "application_id, source_file, parser_version", "24 months", "Store source checksum"],
            ["Screening Result", "screening_id", "score, confidence, rationale, model_version", "24 months", "Immutable after commit"],
            ["Assessment Result", "assessment_id", "type, score, provider_ref, attempt_no", "24 months", "Supports retries and resumes"],
            ["Interview Feedback", "feedback_id", "interviewer_id, rubric_scores, recommendation", "24 months", "Mandatory rubric completeness"],
            ["Final Decision", "decision_id", "outcome, approvers, reason_code", "7 years", "Compliance-critical record"],
            ["Communication Log", "comm_id", "template_id, channel, provider_status", "24 months", "Supports delivery audit"],
            ["Audit Event", "event_id", "actor, action, entity_type, entity_id, timestamp", "7 years", "Append-only"],
        ],
    )

    document.add_heading("7.2 Field-Level Data Dictionary (Representative)", level=2)
    add_table_docx(
        document,
        ["Field", "Type", "Required", "Validation", "Source"],
        [
            ["candidate.email", "string", "Yes", "RFC-compliant format, unique per tenant", "User input"],
            ["candidate.phone", "string", "Conditional", "E.164 format", "User input"],
            ["application.stage", "enum", "Yes", "Must match workflow state machine", "System"],
            ["screening.score", "decimal(5,2)", "Yes", "0 to 100", "AI service"],
            ["screening.confidence", "decimal(5,2)", "Yes", "0 to 1", "AI service"],
            ["decision.outcome", "enum", "Yes", "offer/reject/hold/withdraw", "HR manager"],
            ["decision.reason_code", "string", "Yes for reject", "Taxonomy lookup", "HR manager"],
            ["comm.provider_status", "enum", "Yes", "queued/sent/delivered/bounced/failed", "Email provider"],
        ],
    )

    document.add_heading("7.3 Integration Requirements", level=2)
    add_table_docx(
        document,
        ["Integration", "Direction", "Purpose", "Contract", "Error Handling"],
        [
            ["ATS", "Bi-directional", "Requisition sync and candidate status updates", "API and scheduled sync", "Retry then dead-letter with reconciliation report"],
            ["Assessment Provider", "Bi-directional", "Launch tests and receive scores", "Signed webhook callbacks", "Idempotent ingest and replay queue"],
            ["Email Provider", "Outbound", "Offer and rejection communication", "Template payload API", "Retry and fallback channel"],
            ["Calendar", "Bi-directional", "Interview scheduling", "OAuth API", "Manual scheduling fallback"],
            ["Identity Provider", "Inbound", "SSO for internal roles", "OIDC/SAML", "Fail closed for unauthorized claims"],
        ],
    )

    document.add_heading("7.4 Data Migration and Backfill", level=2)
    add_bullets_docx(
        document,
        [
            "Map legacy candidate and requisition identifiers to new canonical IDs.",
            "Run dry migration in non-production with reconciliation report and null-field audit.",
            "Backfill only active and recently closed requisitions to reduce low-value historical noise.",
            "Keep immutable import manifest including source extract date, checksum, and row counts.",
        ],
    )

    document.add_heading("8. Non-Functional Requirements", level=1)
    add_table_docx(
        document,
        ["Category", "Target", "Measurement Method", "Breach Action"],
        [
            ["API Latency", "P95 < 2 seconds for core read/write endpoints at 500 concurrent users", "APM percentile dashboard", "Trigger performance incident and autoscaling check"],
            ["Screening Throughput", "95 percent of screenings complete within 4 minutes", "Workflow completion histogram", "Activate fallback mode and ops alert"],
            ["Availability", "99.9 percent monthly uptime for core workflow services", "Synthetic probes and uptime monitor", "Incident postmortem and corrective action"],
            ["Security", "Zero critical OWASP issues in pre-release scans", "DAST and SAST reports", "Release block"],
            ["Audit Coverage", "100 percent of stage transitions and decisions logged", "Audit sampling and dashboard", "Compliance escalation"],
            ["Accessibility", "WCAG 2.2 AA for candidate and reviewer critical flows", "Accessibility test suite", "UI release gate"],
            ["Data Integrity", "No orphan application records across stage transitions", "Daily integrity jobs", "Repair workflow and alert"],
            ["Recovery", "RPO <= 15 minutes, RTO <= 2 hours", "Disaster recovery drills", "Executive incident review"],
        ],
    )

    document.add_heading("9. Testing Strategy", level=1)
    document.add_heading("9.1 Test Scope Matrix", level=2)
    add_table_docx(
        document,
        ["Test Type", "Objective", "Minimum Coverage", "Owner"],
        [
            ["Unit", "Validate domain logic and rule enforcement", "85 percent on core domains", "Engineering"],
            ["API Integration", "Validate service contracts and external adapters", "All integration endpoints", "Engineering QA"],
            ["Workflow E2E", "Validate full path from registration to final decision", "All branches and edge states", "QA and Product"],
            ["Performance", "Validate load, stress, and endurance thresholds", "Peak hiring profile", "Performance QA"],
            ["Security", "Validate authz, injection, and data exposure controls", "OWASP top controls", "Security QA"],
            ["Accessibility", "Validate WCAG conformance", "Critical flows and forms", "UX QA"],
            ["UAT", "Validate operational fit and decision quality", "Key user journeys", "HR Ops and Recruiters"],
        ],
    )

    document.add_heading("9.2 Workflow Test Cases (Representative)", level=2)
    add_table_docx(
        document,
        ["Test ID", "Scenario", "Precondition", "Expected Result"],
        [
            ["TC-WF-01", "Candidate registers with valid OTP", "New user", "Account verified and login enabled"],
            ["TC-WF-02", "Duplicate application attempt", "Existing application in cooling period", "System blocks with reason message"],
            ["TC-WF-03", "Low confidence AI result", "Screening confidence below threshold", "Candidate routed to mandatory HR review"],
            ["TC-WF-04", "Fresher branch enforcement", "Candidate classified fresher", "Aptitude required before coding"],
            ["TC-WF-05", "Assessment callback retry", "Transient provider failure", "Idempotent retry and single score record"],
            ["TC-WF-06", "Final decision without HR round", "HR round missing", "Decision action blocked"],
            ["TC-WF-07", "Offer send with missing approval", "Compensation above approval limit", "Offer send blocked and approval prompt shown"],
            ["TC-WF-08", "Rejected candidate communication", "Final outcome reject", "Rejection mail sent and logged"],
        ],
    )

    document.add_heading("9.3 Entry and Exit Criteria", level=2)
    add_bullets_docx(
        document,
        [
            "SIT entry: all critical APIs implemented, test data seeded, and integration stubs available.",
            "SIT exit: zero critical defects, zero high-severity open defects without approved waiver.",
            "UAT entry: signed SIT report, stable build for two consecutive days, role-specific user training complete.",
            "Go-live exit: operational runbooks, on-call roster, and rollback strategy approved.",
        ],
    )

    document.add_heading("9.4 Automation Testing Framework", level=2)
    document.add_paragraph(
        "Automation testing shall be implemented using Playwright with TypeScript for cross-browser, end-to-end validation across the hiring workflow."
    )
    add_table_docx(
        document,
        ["Area", "Automation Approach", "Tooling"],
        [
            ["E2E Candidate Flow", "Automate registration, application, resume upload, and stage progression", "Playwright + TypeScript"],
            ["HR and Recruiter Journeys", "Automate shortlist, routing, decision, and communication trigger paths", "Playwright + TypeScript"],
            ["Regression Suite", "Run smoke and full regression in CI for each release candidate", "Playwright Test Runner with CI integration"],
            ["Cross-Browser Validation", "Validate Chromium, Firefox, and WebKit behavior for critical screens", "Playwright projects configuration"],
            ["Artifacts and Reporting", "Capture traces, screenshots, and videos on failure", "Playwright trace viewer and HTML reports"],
        ],
    )

    document.add_heading("10. Risks and Mitigations", level=1)
    add_table_docx(
        document,
        ["Risk", "Probability", "Impact", "Owner", "Mitigation", "Contingency"],
        [
            ["AI false negative rejects strong candidates", "Medium", "High", "HR Operations", "Human review gate and threshold tuning", "Manual shortlist audit"],
            ["Assessment provider outage", "Medium", "Medium", "Engineering", "Retry queue and status monitor", "Reschedule and alternate test provider"],
            ["Calendar integration drift", "Low", "Medium", "Product Ops", "Timezone and calendar reconciliation checks", "Manual scheduling fallback"],
            ["Unauthorized data access", "Low", "High", "Security Lead", "RBAC hardening and privileged access review", "Incident response playbook"],
            ["Delayed recruiter actions causing SLA breach", "High", "Medium", "HR Manager", "SLA dashboard and escalation alerts", "Temporary queue rebalancing"],
            ["Model version regression", "Medium", "Medium", "Data Science Lead", "Shadow evaluation before promotion", "Rollback to prior model version"],
            ["Email deliverability degradation", "Medium", "Medium", "IT Ops", "Domain authentication and provider monitoring", "Secondary provider failover"],
        ],
    )

    document.add_heading("11. Roles, Permissions, and Access Matrix", level=1)
    add_table_docx(
        document,
        ["Capability", "Applicant", "Recruiter", "HR Reviewer", "Technical Interviewer", "HR Manager", "System Admin"],
        [
            ["Register/Login", "Y", "Y", "Y", "Y", "Y", "Y"],
            ["View Job Listings", "Y", "Y", "Y", "N", "Y", "Y"],
            ["Submit Application", "Y", "N", "N", "N", "N", "N"],
            ["Upload Resume", "Y", "N", "N", "N", "N", "N"],
            ["View AI Screening Result", "N", "Y", "Y", "N", "Y", "Y"],
            ["Shortlist/Reject at HR Stage", "N", "Y", "Y", "N", "Y", "Y"],
            ["Assign Experience Path", "N", "Y", "Y", "N", "Y", "Y"],
            ["Launch Aptitude/Coding", "N", "Y", "Y", "N", "Y", "Y"],
            ["Submit Technical Scorecard", "N", "N", "N", "Y", "N", "Y"],
            ["Run HR Round", "N", "N", "Y", "N", "Y", "Y"],
            ["Publish Final Decision", "N", "N", "N", "N", "Y", "Y"],
            ["Send Offer/Rejection", "N", "Y", "Y", "N", "Y", "Y"],
            ["Edit Communication Templates", "N", "N", "N", "N", "Y", "Y"],
            ["View Audit Logs", "N", "N", "Y", "N", "Y", "Y"],
            ["Manage System Configuration", "N", "N", "N", "N", "N", "Y"],
        ],
    )

    document.add_heading("12. Communication Templates", level=1)
    document.add_heading("12.1 Offer Email Template", level=2)
    document.add_paragraph("Subject: Congratulations, <Candidate Name> - Offer for <Job Title> at <Company Name>")
    document.add_paragraph(
        "Body:\n"
        "Dear <Candidate Name>,\n\n"
        "We are pleased to offer you the position of <Job Title> at <Company Name>. Your profile and interview performance demonstrated strong alignment with our role expectations.\n\n"
        "Offer details:\n"
        "- Department: <Department>\n"
        "- Location: <Location>\n"
        "- Employment Type: <Employment Type>\n"
        "- Proposed Joining Date: <Joining Date>\n"
        "- Compensation: <Compensation Details>\n\n"
        "Please review the attached offer details and confirm your decision by <Response Deadline>.\n\n"
        "Regards,\n"
        "<HR Manager Name>\n"
        "<Company Name>"
    )

    document.add_heading("12.2 Rejection Email Template", level=2)
    document.add_paragraph("Subject: Update on your application for <Job Title> at <Company Name>")
    document.add_paragraph(
        "Body:\n"
        "Dear <Candidate Name>,\n\n"
        "Thank you for participating in our hiring process for <Job Title>. After careful evaluation, we will not be moving forward with your application for this role.\n\n"
        "We appreciate your interest and the effort you invested. We encourage you to apply for future opportunities aligned with your profile.\n\n"
        "Regards,\n"
        "Talent Acquisition Team\n"
        "<Company Name>"
    )

    document.add_heading("12.3 Interview Invitation Template", level=2)
    document.add_paragraph("Subject: Interview Scheduled for <Job Title> - <Company Name>")
    document.add_paragraph(
        "Body:\n"
        "Dear <Candidate Name>,\n\n"
        "Your interview for <Job Title> is scheduled on <Interview Date> at <Interview Time> (<Time Zone>).\n"
        "Mode: <Interview Mode>. Panel: <Interviewer Name>.\n\n"
        "Please confirm your availability by <Confirmation Deadline>.\n\n"
        "Regards,\n"
        "Talent Acquisition Team"
    )

    document.add_heading("12.4 Reminder Template", level=2)
    document.add_paragraph("Subject: Reminder - Action Pending for <Job Title> Application")
    document.add_paragraph(
        "Body:\n"
        "Dear <Candidate Name>,\n\n"
        "This is a reminder that your action is pending for the <Job Title> hiring process.\n"
        "Pending Action: <Action Type>. Deadline: <Action Deadline>.\n\n"
        "If you need support, reply to this email.\n\n"
        "Regards,\n"
        "Talent Acquisition Team"
    )

    document.add_heading("13. Reporting and MIS", level=1)
    add_table_docx(
        document,
        ["Report", "Audience", "Frequency", "KPI Examples"],
        [
            ["Hiring Funnel Dashboard", "HR leadership", "Daily", "Applied, screened, shortlisted, interviewed, offered"],
            ["SLA Breach Report", "HR operations", "Daily", "Review latency, pending approvals, communication delay"],
            ["AI Screening Quality", "Talent strategy", "Weekly", "Precision proxy, confidence distribution, override rate"],
            ["Interviewer Performance", "Engineering leadership", "Weekly", "Scorecard completeness, turnaround time"],
            ["Offer Conversion Report", "Business and HR heads", "Weekly", "Offer acceptance, decline reasons, response time"],
            ["Compliance Audit Pack", "Compliance and security", "Monthly", "Access log summaries, decision trace coverage"],
        ],
    )

    document.add_heading("14. Compliance, Security, and Governance", level=1)
    add_bullets_docx(
        document,
        [
            "Access control must enforce least privilege with role-based permission checks at API and service layers.",
            "PII fields must be encrypted at rest and protected in transit through TLS.",
            "No passwords, tokens, or sensitive candidate data in application logs.",
            "Critical actions (shortlist, reject, final decision, offer) must have immutable audit events.",
            "Data retention and deletion operations must align with legal and policy requirements.",
            "Security scanning (SAST and DAST) is mandatory before each release milestone.",
        ],
    )

    document.add_heading("15. Delivery Roadmap", level=1)
    add_table_docx(
        document,
        ["Milestone", "Primary Scope", "Dependencies", "Definition of Done"],
        [
            ["M1", "Identity, registration, application basics", "IdP and requisition feed", "Candidate can register and apply E2E"],
            ["M2", "Resume parsing, AI screening, HR review", "AI model readiness", "Shortlist and reject path operational"],
            ["M3", "Assessment and interview orchestration", "Assessment and calendar integrations", "Fresher and experienced branches tested"],
            ["M4", "Decision governance and communications", "Approval policy and templates", "Offer and rejection workflows operational"],
            ["M5", "Reporting, compliance hardening, and go-live", "Monitoring stack", "Production readiness sign-off"],
        ],
    )

    document.add_heading("16. Open Issues and Assumptions", level=1)
    add_table_docx(
        document,
        ["Type", "Description", "Owner", "Target Resolution"],
        [
            ["Open Issue", "Define final fresher versus experienced threshold policy", "HR Operations", "Before M2"],
            ["Open Issue", "Confirm compensation approval matrix by business unit", "HR Manager", "Before M4"],
            ["Open Issue", "Select primary and secondary email providers", "IT Operations", "Before M4"],
            ["Assumption", "Assessment provider supports signed webhook callbacks", "Engineering", "Validate in M2"],
            ["Assumption", "ATS can receive near real-time status updates", "Product", "Validate in M1"],
            ["Assumption", "Recruiters can meet review SLAs with dashboard alerts", "HR Leadership", "Validate during UAT"],
        ],
    )

    document.add_heading("17. Glossary", level=1)
    add_table_docx(
        document,
        ["Term", "Definition"],
        [
            ["ATS", "Applicant Tracking System used for requisition and candidate status integration"],
            ["SLA", "Service Level Agreement for expected turnaround time at each stage"],
            ["RBAC", "Role-Based Access Control"],
            ["RPO", "Recovery Point Objective"],
            ["RTO", "Recovery Time Objective"],
            ["Idempotency", "Guarantee that repeated processing of same request has same effect"],
            ["Confidence", "Model certainty signal attached to AI screening output"],
            ["Decision Trace", "Complete historical evidence of how final hiring decision was reached"],
        ],
    )

    document.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleCenter",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SubCenter",
            parent=styles["Normal"],
            alignment=TA_CENTER,
            fontName="Helvetica-Oblique",
            fontSize=10,
            leading=12,
            textColor=colors.HexColor("#475569"),
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=17,
            spaceBefore=9,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=11.3,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Small",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=9.5,
            spaceAfter=2,
        )
    )
    return styles


def pdf_table(headers, rows, widths, styles):
    data = [[Paragraph(str(header), styles["Small"]) for header in headers]]
    for row in rows:
        data.append([Paragraph(str(value), styles["Small"]) for value in row])
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def pdf_bullets(items, styles):
    flow_items = []
    for item in items:
        flow_items.append(ListItem(Paragraph(item, styles["Body"]), leftIndent=9))
    return ListFlowable(flow_items, bulletType="bullet", start="circle")


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=0.56 * inch,
        rightMargin=0.56 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    story = []

    story.append(Paragraph("AI Interview Application", styles["TitleCenter"]))
    story.append(Paragraph("Business Requirements Document", styles["TitleCenter"]))
    story.append(
        Paragraph(
            "Architect perspective with implementation-ready detail and workflow governance",
            styles["SubCenter"],
        )
    )

    story.append(
        pdf_table(
            ["Field", "Value"],
            [
                ["Version", "2.0"],
                ["Prepared Date", date.today().isoformat()],
                ["Document Type", "Workflow-based enterprise BRD"],
                ["Source Workflow", "Registration to Offer/Rejection with fresher and experienced branches"],
                ["Status", "Draft for architecture, product, and operations sign-off"],
            ],
            [1.8 * inch, 5.0 * inch],
            styles,
        )
    )
    story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph("1. Executive Summary", styles["Section"]))
    story.append(
        Paragraph(
            "The platform standardizes hiring execution from account registration to final offer or rejection. It combines AI-assisted screening with mandatory human gates, reducing manual delays while retaining decision accountability.",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "Architecture is modular and policy-driven: deterministic stage transitions for critical decisions, event-driven processing for long-running tasks, and immutable audit trails for compliance-sensitive actions.",
            styles["Body"],
        )
    )

    story.append(Paragraph("2. Detailed Workflow", styles["Section"]))
    story.append(
        pdf_table(
            ["Step", "Actor", "Trigger", "Output"],
            [[row[1], row[2], row[3], row[5]] for row in workflow_steps()],
            [1.7 * inch, 1.1 * inch, 2.2 * inch, 2.0 * inch],
            styles,
        )
    )

    story.append(Paragraph("3. Functional Requirement Modules", styles["Section"]))
    for module_name, rows in functional_modules():
        story.append(Paragraph(module_name, styles["Body"]))
        story.append(
            pdf_table(
                ["ID", "Requirement", "Priority", "Acceptance"],
                [[r[0], r[1], r[2], r[4]] for r in rows],
                [0.9 * inch, 3.0 * inch, 0.7 * inch, 2.4 * inch],
                styles,
            )
        )

    story.append(Paragraph("4. Business Rules", styles["Section"]))
    story.append(
        pdf_table(
            ["Rule", "If", "Then"],
            [
                ["BR-01", "Resume missing", "Block screening"],
                ["BR-02", "Low AI confidence", "Force HR review"],
                ["BR-04", "Fresher", "Aptitude before coding"],
                ["BR-06", "Missing mandatory stage", "Block final decision"],
                ["BR-07", "High compensation", "Require extra approval"],
                ["BR-10", "Email bounce", "Create recruiter task"],
                ["BR-15", "Deletion request", "Anonymize with legal retention"],
            ],
            [0.8 * inch, 2.8 * inch, 3.4 * inch],
            styles,
        )
    )

    story.append(Paragraph("5. UI and UX", styles["Section"]))
    story.append(
        pdf_bullets(
            [
                "Dashboard-first workflow with explicit next action and SLA indicator.",
                "Accessible forms with keyboard support and readable error messaging.",
                "Candidate forms use autosave and progress indicators to reduce drop-off.",
                "No hidden status changes without explicit reviewer confirmation.",
            ],
            styles,
        )
    )

    story.append(Paragraph("6. Architecture", styles["Section"]))
    story.append(
        pdf_table(
            ["Layer", "Technology", "Purpose"],
            [
                ["Frontend", "React", "Candidate and recruiter-facing web interfaces"],
                ["Backend", "Node.js", "Core APIs and workflow orchestration"],
                ["AI Integration", "Python", "AI screening and model integration services"],
                ["Database", "PostgreSQL", "Primary relational data store"],
            ],
            [1.5 * inch, 1.3 * inch, 3.2 * inch],
            styles,
        )
    )
    story.append(
        pdf_table(
            ["Component", "Purpose", "Constraint"],
            [
                ["Workflow Orchestrator", "State transition control", "Deterministic guards"],
                ["Resume and AI Service", "Parse and score", "Versioned model outputs"],
                ["Assessment Adapter", "External test sync", "Idempotent callbacks"],
                ["Decision Service", "Final approval logic", "Approval matrix enforcement"],
                ["Audit Service", "Immutable events", "Append-only storage"],
            ],
            [1.8 * inch, 2.4 * inch, 2.8 * inch],
            styles,
        )
    )

    story.append(Paragraph("7. Data and Integrations", styles["Section"]))
    story.append(
        pdf_table(
            ["Area", "Details"],
            [
                ["Core Entities", "Candidate, Application, Resume, Screening, Assessment, Feedback, Decision, Communication, Audit"],
                ["ATS Integration", "Requisition sync plus candidate status updates"],
                ["Assessment Integration", "Launch tests and ingest signed callbacks"],
                ["Email Integration", "Tokenized templates with provider status tracking"],
                ["Retention", "Decision and audit logs retained for compliance horizon"],
            ],
            [1.8 * inch, 5.0 * inch],
            styles,
        )
    )

    story.append(Paragraph("8. Non-Functional Targets", styles["Section"]))
    story.append(
        pdf_table(
            ["Category", "Target"],
            [
                ["Latency", "P95 < 2 seconds on core APIs under expected load"],
                ["Availability", "99.9 percent monthly uptime"],
                ["Audit", "100 percent decision path traceability"],
                ["Security", "No critical OWASP findings at release"],
                ["Recovery", "RPO <= 15 min, RTO <= 2 hours"],
            ],
            [1.8 * inch, 5.0 * inch],
            styles,
        )
    )

    story.append(Paragraph("9. Testing Strategy", styles["Section"]))
    story.append(
        pdf_table(
            ["Type", "Scope", "Owner"],
            [
                ["Unit", "Business rules and stage guards", "Engineering"],
                ["Integration", "ATS, assessment, email, calendar", "Engineering QA"],
                ["E2E", "Both candidate branches and edge states", "QA and Product"],
                ["Security", "Authz and data protection controls", "Security QA"],
                ["UAT", "Operational readiness and usability", "HR Ops"],
            ],
            [1.4 * inch, 3.8 * inch, 1.6 * inch],
            styles,
        )
    )
    story.append(
        Paragraph(
            "Automation testing framework: Playwright with TypeScript for cross-browser E2E validation, regression, and CI execution.",
            styles["Body"],
        )
    )

    story.append(Paragraph("10. Risks", styles["Section"]))
    story.append(
        pdf_table(
            ["Risk", "Mitigation"],
            [
                ["AI misclassification", "Human review gate and threshold tuning"],
                ["Provider outage", "Retry queue and manual fallback"],
                ["Data access incident", "RBAC hardening and incident playbook"],
                ["SLA breaches", "Escalation dashboards and queue balancing"],
            ],
            [2.3 * inch, 4.5 * inch],
            styles,
        )
    )

    story.append(Paragraph("11. Roles and Access", styles["Section"]))
    story.append(
        pdf_table(
            ["Role", "Core Responsibilities"],
            [
                ["Applicant", "Register, apply, upload documents, complete assessments"],
                ["Recruiter", "Manage funnel, shortlist operations, communications"],
                ["HR Reviewer", "Validate screening outputs and progression decisions"],
                ["Technical Interviewer", "Run interviews and submit rubric scorecards"],
                ["HR Manager", "Finalize decision and approvals"],
                ["System Admin", "Configuration, access controls, integration health"],
            ],
            [1.6 * inch, 5.2 * inch],
            styles,
        )
    )

    story.append(Paragraph("12. Communication Templates", styles["Section"]))
    story.append(Paragraph("Includes offer, rejection, interview invitation, and reminder templates with tokenized placeholders.", styles["Body"]))

    story.append(Paragraph("13. Compliance and Governance", styles["Section"]))
    story.append(
        pdf_bullets(
            [
                "Least-privilege access with API-level authorization checks.",
                "Encrypted PII storage and secure transport.",
                "Immutable audit records for shortlist and final decision actions.",
                "Retention and deletion aligned to legal and policy requirements.",
            ],
            styles,
        )
    )

    story.append(Paragraph("14. Delivery Roadmap", styles["Section"]))
    story.append(
        pdf_table(
            ["Milestone", "Outcome"],
            [
                ["M1", "Identity and application flow live in SIT"],
                ["M2", "Screening and HR review production-ready"],
                ["M3", "Interview orchestration and decision controls complete"],
                ["M4", "Communications, reporting, and go-live readiness signed off"],
            ],
            [1.1 * inch, 5.7 * inch],
            styles,
        )
    )

    story.append(Paragraph("15. Open Issues and Glossary", styles["Section"]))
    story.append(Paragraph("Open items include path classification policy, compensation approval thresholds, and provider finalization.", styles["Body"]))
    story.append(Paragraph("Glossary covers ATS, SLA, RBAC, RPO, RTO, and idempotency terms for implementation alignment.", styles["Body"]))

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#475569"))
        canvas.drawRightString(A4[0] - 0.56 * inch, 0.38 * inch, f"Page {document.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    build_docx()
    build_pdf()
